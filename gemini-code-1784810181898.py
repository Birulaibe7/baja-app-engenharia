import streamlit as st
import pandas as pd
from datetime import datetime
import os
import tempfile
import io
from PIL import Image
from docx import Document
from docx.shared import Inches
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo

# ==========================================
# CONFIGURAÇÕES DA PÁGINA E BANCO DE DADOS
# ==========================================
st.set_page_config(page_title="Central de Engenharia - Baja", page_icon="⚙️", layout="wide")

EXCEL_FALHAS = "historico_baja.xlsx"
STATUS_LIST = ["🔴 Ocorrência Registrada", "🟡 Em Análise", "🔵 Fila de Usinagem", "🟢 Validado / Fechado"]
COLUNAS = [
    "ID", "Status", "Data_Ocorrencia", "Componente", "Area", "Responsavel", 
    "Contexto", "Piloto", "Atividade", "Sintomas", "Inspecao", "Drive_Link", 
    "Porque1", "Porque2", "Porque3", "Porque4", "Porque5", 
    "Acao", "Modificacoes", "Resultado", "Parecer"
]

def inicializar_bd():
    if not os.path.exists(EXCEL_FALHAS):
        df = pd.DataFrame(columns=COLUNAS)
        df.to_excel(EXCEL_FALHAS, index=False)

inicializar_bd()

def carregar_dados():
    return pd.read_excel(EXCEL_FALHAS)

def salvar_dados(df):
    df.to_excel(EXCEL_FALHAS, index=False)

# ==========================================
# FUNÇÕES DE GERADORES DE RELATÓRIO
# ==========================================
def processar_imagem_temp(uploaded_file):
    if uploaded_file is None:
        return None
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        img = Image.open(uploaded_file)
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        img.save(tmp.name)
        return tmp.name

def limpar_texto_pdf(texto):
    """Remove emojis e acentos pesados para o FPDF não dar crash"""
    if not isinstance(texto, str):
        return str(texto)
    texto_sem_emoji = texto.replace("🔴 ", "").replace("🟡 ", "").replace("🔵 ", "").replace("🟢 ", "")
    texto_limpo = unicodedata.normalize('NFKD', texto_sem_emoji).encode('ASCII', 'ignore').decode('ASCII')
    return texto_limpo

def gerar_word(dados, img_paths):
    doc = Document()
    doc.add_heading(f"RAF - RELATÓRIO DE ANÁLISE DE FALHA ({dados['ID']})", level=1)
    
    doc.add_heading("1. Identificação", level=2)
    doc.add_paragraph(f"Status Atual: {dados['Status']}")
    doc.add_paragraph(f"Componente: {dados['Componente']} | Área: {dados['Area']}")
    doc.add_paragraph(f"Data: {dados['Data_Ocorrencia']} | Responsável: {dados['Responsavel']}")
    doc.add_paragraph(f"Link CAD (Drive): {dados['Drive_Link']}")

    doc.add_heading("2. Registro do Fato", level=2)
    doc.add_paragraph(f"Contexto: {dados['Contexto']}")
    doc.add_paragraph(f"Piloto: {dados['Piloto']} | Atividade: {dados['Atividade']}")
    doc.add_paragraph(f"Sintomas: {dados['Sintomas']}")
    if img_paths.get('quebra'):
        doc.add_picture(img_paths['quebra'], width=Inches(2.5)) 

    doc.add_heading("3. Análise da Causa Raiz", level=2)
    doc.add_paragraph(f"Inspeção Visual: {dados['Inspecao']}")
    for i in range(1, 6):
        doc.add_paragraph(f"{i}. {dados[f'Porque{i}']}")

    doc.add_heading("4. Solução e Validação", level=2)
    doc.add_paragraph(f"Ação Escolhida: {dados['Acao']}")
    doc.add_paragraph(f"Modificações: {dados['Modificacoes']}")
    if img_paths.get('nova'):
        doc.add_picture(img_paths['nova'], width=Inches(2.5)) 
    
    doc.add_paragraph(f"Resultado: {dados['Resultado']}")
    doc.add_paragraph(f"Parecer Final: {dados['Parecer']}")

    tmp_path = tempfile.mktemp(suffix=".docx")
    doc.save(tmp_path)
    return tmp_path

    pdf.ln(5)
    add_linha("Status", dados['Status'])
    add_linha("Componente", dados['Componente'])
    add_linha("Data", dados['Data_Ocorrencia'])
    add_linha("Link CAD", dados['Drive_Link'])
    
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "2. O Fato", ln=True)
    add_linha("Contexto", dados['Contexto'])
    if img_paths.get('quebra'):
        pdf.image(img_paths['quebra'], w=60)
        pdf.ln(2)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "3. Analise de Causa Raiz", ln=True)
    add_linha("Inspecao Visual", dados['Inspecao'])
    add_linha("Causa Raiz", dados['Porque5'])

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "4. Solucao", ln=True)
    add_linha("Acao", dados['Acao'])
    if img_paths.get('nova'):
        pdf.image(img_paths['nova'], w=60)
        pdf.ln(2)
    add_linha("Parecer Final", dados['Parecer'])

    tmp_path = tempfile.mktemp(suffix=".pdf")
    pdf.output(tmp_path)
    return tmp_path

# ==========================================
# INTERFACE DO USUÁRIO (ABAS)
# ==========================================
st.title("⚙️ Central de Engenharia & Oficina - Baja")

aba1, aba2, aba3 = st.tabs(["📝 Kanban (Lançar/Editar RAF)", "🗜️ Fila de Usinagem", "🗂️ Exportar e Relatórios"])

# --- ABA 1: KANBAN E EDIÇÃO CONTÍNUA ---
with aba1:
    df = carregar_dados()
    # Converte tudo para texto de forma segura para o Pandas não dar chilique
    df = df.fillna("").astype(str)
    df = df.replace("nan", "")
    
    opcoes_raf = ["✨ Abrir NOVO Registro"] + df["ID"].tolist()
    raf_selecionado = st.selectbox("Selecione a Ação:", opcoes_raf)

    if raf_selecionado != "✨ Abrir NOVO Registro":
        dados_atuais = df[df["ID"] == raf_selecionado].iloc[0].to_dict()
        st.info(f"Editando o registro: **{raf_selecionado}**")
    else:
        novo_id = f"RAF-{datetime.now().year}-{str(len(df)+1).zfill(3)}"
        dados_atuais = {col: "" for col in COLUNAS}
        dados_atuais["ID"] = novo_id
        dados_atuais["Status"] = STATUS_LIST[0]
        st.success(f"Criando novo registro: **{novo_id}**")

    # O parâmetro clear_on_submit=True zera o formulário após salvar!
    with st.form("form_kanban", clear_on_submit=True):
        col1, col2 = st.columns(2)
        status_idx = STATUS_LIST.index(dados_atuais["Status"]) if dados_atuais["Status"] in STATUS_LIST else 0
        
        status_novo = col1.selectbox("Status do RAF (Kanban)", STATUS_LIST, index=status_idx)
        drive_link = col2.text_input("🔗 Link da Pasta no Drive (CAD 3D/2D)", value=str(dados_atuais["Drive_Link"]))
        
        st.divider()
        st.subheader("1. Identificação e Ocorrência")
        c1, c2, c3 = st.columns(3)
        comp = c1.text_input("Componente", value=str(dados_atuais["Componente"]))
        area = c2.selectbox("Área", ["Suspensão", "Powertrain", "Chassi", "Freios", "Elétrica"], index=0)
        data = c3.date_input("Data da Ocorrência")
        
        c4, c5, c6 = st.columns(3)
        resp = c4.text_input("Responsável Técnico", value=str(dados_atuais["Responsavel"]))
        piloto = c5.text_input("Piloto / Operador", value=str(dados_atuais["Piloto"]))
        atividade = c6.selectbox("Tipo de Atividade", ["Treino/Oficina", "Competição", "Teste Dinâmico"], index=0)
        
        contexto = st.text_area("Contexto Detalhado da Quebra", value=str(dados_atuais["Contexto"]))
        sintomas = st.text_area("Sintomas da Falha", value=str(dados_atuais["Sintomas"]))
        
        st.divider()
        st.subheader("2. Análise da Causa Raiz (Os 5 Porquês)")
        insp = st.text_input("Inspeção Macroscópica da Superfície", value=str(dados_atuais["Inspecao"]))
        
        pq1 = st.text_input("1. Por que falhou mecanicamente?", value=str(dados_atuais["Porque1"]))
        pq2 = st.text_input("2. Por que ocorreu essa condição?", value=str(dados_atuais["Porque2"]))
        pq3 = st.text_input("3. Por que o esforço agiu assim?", value=str(dados_atuais["Porque3"]))
        pq4 = st.text_input("4. Por que o projeto não mitigou?", value=str(dados_atuais["Porque4"]))
        pq5 = st.text_input("5. Por que não foi previsto? (Causa Raiz)", value=str(dados_atuais["Porque5"]))
        
        st.divider()
        st.subheader("3. Solução e Validação")
        acao = st.text_area("Ação Técnica Escolhida & Justificativa", value=str(dados_atuais["Acao"]))
        mods = st.text_area("Modificações de Material/Usinagem", value=str(dados_atuais["Modificacoes"]))
        
        # Mantemos apenas um campo descritivo e de status final consolidado
        parecer_opcoes = ["Em Aberto / Em Teste", "APROVADO PARA COMPETIÇÃO", "REPROVADO - NECESSITA REVISÃO"]
        parecer_atual = dados_atuais["Parecer"] if dados_atuais["Parecer"] in parecer_opcoes else parecer_opcoes[0]
        resultado = st.selectbox("Parecer e Resultado dos Testes de Campo", parecer_opcoes, index=parecer_opcoes.index(parecer_atual))
        
        st.info("💡 As fotos devem ser anexadas na Aba 3 na hora de gerar o Relatório Final em Word.")
        submit = st.form_submit_button("Salvar no Banco de Dados", use_container_width=True)

    if submit:
        nova_linha = dados_atuais.copy()
        nova_linha.update({
            "Status": status_novo, "Drive_Link": drive_link, "Componente": comp, 
            "Area": area, "Data_Ocorrencia": data.strftime("%d/%m/%Y"), "Responsavel": resp, 
            "Piloto": piloto, "Atividade": atividade, "Sintomas": sintomas,
            "Contexto": contexto, "Inspecao": insp, 
            "Porque1": pq1, "Porque2": pq2, "Porque3": pq3, "Porque4": pq4, "Porque5": pq5,
            "Acao": acao, "Modificacoes": mods, "Resultado": resultado, "Parecer": resultado
        })

        if raf_selecionado == "✨ Abrir NOVO Registro":
            df = pd.concat([df, pd.DataFrame([nova_linha])], ignore_index=True)
        else:
            idx_update = df[df["ID"] == raf_selecionado].index[0]
            for coluna, valor in nova_linha.items():
                df.at[idx_update, coluna] = valor
        
        salvar_dados(df)
        st.success("✅ Dados salvos com sucesso! A base principal foi atualizada.")
        st.rerun()
# --- ABA 2: FILA DE USINAGEM E HISTÓRICO ---
with aba2:
    st.subheader("🗜️ Fila de Produção / Manufatura")
    df = carregar_dados()
    fila = df[df["Status"] == "🔵 Fila de Usinagem"]
    
    if fila.empty:
        st.success("A fila de usinagem está vazia! Nenhuma peça pendente.")
    else:
        for idx, row in fila.iterrows():
            with st.expander(f"🛠️ {row['Componente']} (RAF: {row['ID']})"):
                st.write(f"**Área:** {row['Area']} | **Responsável:** {row['Responsavel']}")
                st.write(f"**O que fazer (Ação):** {row['Acao']}")
                st.markdown(f"🔗 [Acessar Desenho no Google Drive]({row['Drive_Link']})")
                
                if st.button(f"✅ Marcar '{row['Componente']}' como Validado/Usinado", key=f"btn_{idx}"):
                    df.at[idx, "Status"] = "🟢 Validado / Fechado"
                    salvar_dados(df)
                    st.rerun()

    st.divider()
    st.subheader("🏁 Histórico de Peças Concluídas (Validadas)")
    concluidos = df[df["Status"] == "🟢 Validado / Fechado"]
    
    if concluidos.empty:
        st.info("Nenhuma RAF foi finalizada ainda.")
    else:
        st.dataframe(
            concluidos[["ID", "Componente", "Area", "Responsavel", "Data_Ocorrencia"]], 
            hide_index=True, 
            use_container_width=True
        )

# --- ABA 3: GERADOR DE RELATÓRIOS E EXPORTAÇÃO ---
with aba3:
    st.subheader("📄 Gerar Documento Final do RAF")
    df = carregar_dados()
    
    raf_gerar = st.selectbox("Escolha o RAF para gerar documento:", df["ID"].tolist())
    
    col_img1, col_img2 = st.columns(2)
    foto_quebra = col_img1.file_uploader("Foto da Quebra", type=["jpg", "png", "jpeg"])
    foto_nova = col_img2.file_uploader("Foto da Nova Peça", type=["jpg", "png", "jpeg"])
    
    if st.button("📝 Gerar Relatório (Word .docx)", use_container_width=True):
        dados = df[df["ID"] == raf_gerar].iloc[0].to_dict()
        paths = {'quebra': processar_imagem_temp(foto_quebra), 'nova': processar_imagem_temp(foto_nova)}
        word_path = gerar_word(dados, paths)
        with open(word_path, "rb") as f:
            st.download_button("📥 Baixar Relatório Word Pronto", f, file_name=f"{raf_gerar}.docx", use_container_width=True)

    st.divider()
    st.subheader("📊 Exportar Base de Dados (Tabela Dinâmica)")
    st.info("Baixe a planilha oficial. Ela já vem formatada com filtros para você ordenar, buscar e analisar as quebras.")
    
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Base_Dashboard')
        worksheet = writer.sheets['Base_Dashboard']
        
        # 1. Definindo o tamanho total da nossa base de dados
        max_row = worksheet.max_row
        max_col = worksheet.max_column
        col_letter = get_column_letter(max_col)
        tabela_range = f"A1:{col_letter}{max_row}"
        
        # 2. Criando o objeto "Tabela" oficial do Excel
        tabela = Table(displayName="HistoricoFalhas", ref=tabela_range)
        
        # 3. Aplicando um estilo padrão bonitão do Excel (estilo azul com listras e filtros ativos)
        estilo = TableStyleInfo(
            name="TableStyleMedium9", showFirstColumn=False,
            showLastColumn=False, showRowStripes=True, showColumnStripes=False
        )
        tabela.tableStyleInfo = estilo
        worksheet.add_table(tabela)
        
        # 4. Ajustando a largura das colunas e ativando a quebra de texto
        for col_num, column_title in enumerate(df.columns, 1):
            col_let = get_column_letter(col_num)
            if column_title in ["Contexto", "Sintomas", "Inspecao", "Acao", "Porque1", "Porque2", "Porque3", "Porque4", "Porque5"]:
                worksheet.column_dimensions[col_let].width = 45 
            elif column_title in ["ID", "Data_Ocorrencia", "Status", "Responsavel", "Area"]:
                worksheet.column_dimensions[col_let].width = 22
            else:
                worksheet.column_dimensions[col_let].width = 18

        for row in worksheet.iter_rows(min_row=2, max_row=max_row, min_col=1, max_col=max_col):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    st.download_button(
        label="📥 Baixar Planilha Interativa (.xlsx)",
        data=buffer.getvalue(),
        file_name="base_dashboard_baja.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
